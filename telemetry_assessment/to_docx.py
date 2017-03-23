
#
# Author: Abel Corver
#         abel.corver@gmail.com
#         (Anthony Leonardo Lab)
# 
# Last modified: December, 2016
#

# ===========================================================================================
# Imports
# ===========================================================================================

import os, datetime, sys
import comtypes.client
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import process_performance

os.chdir(os.path.dirname(os.path.abspath(__file__)))

def run(fnames):

    # =======================================================================================
    # Settings
    # =======================================================================================

    mdate = datetime.datetime.strftime(datetime.datetime.fromtimestamp( \
        os.path.getmtime(fnames[0][0])), '%Y-%m-%d %H:%M')
    fname = 'telemetry_performance_' + mdate.replace(' ','_').replace(':','-') + '.docx'
    absfname = os.path.abspath( fname )

    # =======================================================================================
    # Create PDF pages
    # =======================================================================================

    document = Document()

    p = document.add_paragraph()
    r = p.add_run()
    r.add_break()
    r.add_break()
    r.add_break()
    r.add_break()
    r.add_text('Telemetry Performance Assessment')
    r.add_break()
    r.add_break()
    r.add_text('Data Files Appear to have been Collected on: ')
    r.add_break()
    r.add_break()
    r.add_text(mdate)
    r.add_break()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Inches(1)
    
    for img in [x for x in os.listdir('.') if '.png' in x]:
        r.add_picture(img, width=Inches(8.5))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

    # Save
    document.save(fname)

    # =======================================================================================
    # Convert to PDF
    # =======================================================================================

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(absfname)
    doc.SaveAs(absfname.replace('.docx','.pdf'), FileFormat=17)
    doc.Close()
    word.Quit()

# ===========================================================================================
# Main entry point
# ===========================================================================================

if __name__ == "__main__":
    run(process_performance.getFilenames())